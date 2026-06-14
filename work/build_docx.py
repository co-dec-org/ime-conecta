import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "work" / "content.json"
OUTPUT_PATH = (
    ROOT
    / "public"
    / "docs"
    / "Informe_IME_Conecta_Gobernanza_Operacional_actualizado.docx"
)

BLUE = RGBColor(19, 112, 239)
CYAN = RGBColor(21, 233, 205)
DARK = RGBColor(8, 8, 8)
INK = RGBColor(46, 45, 50)
MUTED = RGBColor(77, 76, 84)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
BORDER = "B7C2D0"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
      tbl_grid = OxmlElement("w:tblGrid")
      table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.333):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def paragraph(text="", style=None, before=0, after=8, line=1.333, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    set_paragraph_spacing(p, before, after, line)
    if align is not None:
        p.alignment = align
    return p


def add_body(text, bold_prefix=None):
    p = paragraph(after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=INK)
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208


def clean_joined_item(item):
    return str(item).strip().rstrip(".;:")


def add_callout(text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="9ADFD6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2, line=1.208)
    run = p.add_run(text)
    set_run_font(run, color=INK, bold=True)
    paragraph(after=4)


def add_table(columns, rows, widths=None):
    if not rows:
        return
    count = len(columns)
    if widths is None:
        if count == 2:
            widths = [3200, 6160]
        elif count == 3:
            widths = [2100, 3600, 3660]
        elif count == 4:
            widths = [1700, 2300, 3300, 2060]
        else:
            width = 9360 // count
            widths = [width] * count
            widths[-1] += 9360 - sum(widths)

    table = doc.add_table(rows=1, cols=count)
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_table_borders(table)
    header = table.rows[0].cells
    for index, column in enumerate(columns):
        set_cell_shading(header[index], TABLE_FILL)
        p = header[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(column)
        set_run_font(run, bold=True, color=INK, size=9.5)

    for row in rows:
        cells = table.add_row().cells
        for index in range(count):
            value = row[index] if index < len(row) else ""
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, color=INK, size=9.2)
    set_table_width(table, widths)
    paragraph(after=6)


def configure_styles(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.text = "IME Conecta | Gobernanza operacional para una nueva escala gremial"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8.5)
    footer.runs[0].font.color.rgb = MUTED


def add_cover(content):
    paragraph("Propuesta para Directorio IME Chile", after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = paragraph(after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    title_run = title.add_run("Gobernanza operacional para una nueva escala gremial")
    set_run_font(title_run, size=24, color=INK, bold=True)

    subtitle = paragraph(after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle_run = subtitle.add_run(
        "IME Conecta como infraestructura interna para ordenar, profesionalizar y proyectar la gestion de IME Chile."
    )
    set_run_font(subtitle_run, size=13.5, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, [2200, 7160])
    set_table_borders(meta)
    rows = [
        ("Autor", "Patricio Gonzalez Cruz, Director y Tesorero IME Chile"),
        ("Formato", "Informe DOCX/PDF y presentacion PPTX editable"),
        ("Proyecto", "IME Conecta"),
        ("Fecha", "Junio 2026"),
    ]
    for row_index, (label, value) in enumerate(rows):
        cells = meta.rows[row_index].cells
        set_cell_shading(cells[0], TABLE_FILL)
        p0 = cells[0].paragraphs[0]
        p0.add_run(label)
        set_run_font(p0.runs[0], bold=True, color=INK)
        p1 = cells[1].paragraphs[0]
        p1.add_run(value)
        set_run_font(p1.runs[0], color=INK)

    paragraph(after=12)
    add_callout(
        "Tesis: a mayor envergadura gremial, mayor necesidad de estructura operacional, trazabilidad, roles claros y gobernanza de datos."
    )
    doc.add_page_break()


def add_section(section, number):
    add_heading(f"{number:02d}. {section['title']}", level=1)
    if section.get("subtitle"):
        p = paragraph(after=6)
        run = p.add_run(section["subtitle"])
        set_run_font(run, size=12, color=MUTED, italic=True)
    if section.get("body"):
        add_body(section["body"])
    if section.get("definition"):
        add_body(section["definition"])
    if section.get("highlight"):
        add_callout(section["highlight"])
    if section.get("bullets"):
        add_bullets(section["bullets"])
    if section.get("cards"):
        add_bullets(section["cards"])
    if section.get("blocks"):
        add_heading("Componentes", level=2)
        add_bullets(section["blocks"])
    if section.get("phrases"):
        for phrase in section["phrases"]:
            add_callout(phrase)
    if section.get("table"):
        add_table(section["table"]["columns"], section["table"]["rows"])
    if section.get("comparison"):
        comparison = section["comparison"]
        add_table(
            [comparison["leftTitle"], comparison["rightTitle"]],
            comparison["rows"],
            widths=[4680, 4680],
        )
    if section.get("flow"):
        add_heading("Flujo propuesto", level=2)
        add_numbered(section["flow"])
    if section.get("timeline"):
        add_heading("Hitos", level=2)
        rows = [
            [item["label"], "; ".join(clean_joined_item(action) for action in item["items"]) + "."]
            for item in section["timeline"]
        ]
        add_table(["Horizonte", "Acciones"], rows, widths=[2000, 7360])
    if section.get("questions"):
        add_heading("Preguntas para deliberacion", level=2)
        add_numbered(section["questions"])
    if section.get("agreements"):
        add_heading("Materias posibles de acuerdo", level=2)
        add_numbered(section["agreements"])


if __name__ == "__main__":
    with CONTENT_PATH.open("r", encoding="utf8") as handle:
        content = json.load(handle)

    doc = Document()
    configure_styles(doc)
    add_cover(content)

    add_heading("Resumen ejecutivo", level=1)
    for paragraph_text in content["executiveSummary"].split("\n\n"):
        add_body(paragraph_text)

    doc.add_page_break()
    add_heading("Indice de contenidos", level=1)
    for index, section in enumerate(content["sections"], start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(section["title"])
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()
    for index, section in enumerate(content["sections"], start=1):
        if section["id"] in {"portada", "indice"}:
            continue
        add_section(section, index)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
